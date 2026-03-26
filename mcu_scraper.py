"""
銘傳大學選課資訊爬蟲
MCU Course Selection Information Scraper

爬取來源:
- 教務處公告: https://academic.mcu.edu.tw/
- 新生選課說明: https://freshman.mcu.edu.tw/a01/a01-04/
- 英語教學中心選課: https://elc.mcu.edu.tw/選課/
- 課程架構: https://academic.mcu.edu.tw/coursestructure/
"""

import os
import re
import json
import time
import hashlib
import logging
import requests
import pdfplumber
import docx
from datetime import datetime
from pathlib import Path
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup

# ── 設定 ──────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("scraper.log", encoding="utf-8"),
    ],
)
log = logging.getLogger(__name__)

OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "zh-TW,zh;q=0.9,en;q=0.8",
}

SESSION = requests.Session()
SESSION.headers.update(HEADERS)

# 爬取目標清單
TARGET_PAGES = [
    {
        "name": "教務處最新公告",
        "url": "https://academic.mcu.edu.tw/",
        "tag": "announcement",
    },
    {
        "name": "教務處公告列表",
        "url": "https://academic.mcu.edu.tw/?cat=2",
        "tag": "announcement",
    },
    {
        "name": "新生選課注意事項",
        "url": "https://freshman.mcu.edu.tw/a01/a01-04/",
        "tag": "freshman",
    },
    {
        "name": "英語教學中心選課",
        "url": "https://elc.mcu.edu.tw/%E9%81%B8%E8%AA%B2/",
        "tag": "english",
    },
    {
        "name": "開課與師資資訊",
        "url": "https://academic.mcu.edu.tw/coursestructure/",
        "tag": "course_structure",
    },
]

# 選課相關關鍵字（用來篩選有價值的連結）
COURSE_KEYWORDS = [
    "選課", "course", "課程", "開課", "停課", "加退選",
    "選課辦法", "選課時間", "選課注意", "課表", "學分",
    "暑修", "補修", "重修", "校際", "網路選課",
]


# ── 工具函數 ──────────────────────────────────────

def safe_get(url: str, timeout: int = 20) -> requests.Response | None:
    """帶重試的 GET 請求"""
    for attempt in range(3):
        try:
            resp = SESSION.get(url, timeout=timeout, allow_redirects=True)
            resp.raise_for_status()
            return resp
        except requests.RequestException as e:
            log.warning(f"[嘗試 {attempt+1}/3] 無法取得 {url}: {e}")
            time.sleep(2 ** attempt)
    return None


def url_to_filename(url: str) -> str:
    """把 URL 轉成安全的檔案名稱"""
    h = hashlib.md5(url.encode()).hexdigest()[:8]
    parsed = urlparse(url)
    name = re.sub(r"[^\w\-]", "_", parsed.path)[-40:]
    return f"{name}_{h}"


def is_course_related(text: str) -> bool:
    """判斷文字是否與選課相關"""
    text_lower = text.lower()
    return any(kw in text_lower for kw in COURSE_KEYWORDS)


def extract_text_from_pdf(pdf_bytes: bytes, filename: str) -> str:
    """從 PDF bytes 擷取純文字"""
    tmp_path = OUTPUT_DIR / f"_tmp_{filename}.pdf"
    tmp_path.write_bytes(pdf_bytes)
    text_parts = []
    try:
        with pdfplumber.open(str(tmp_path)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(f"--- 第 {i} 頁 ---\n{page_text}")
        text = "\n\n".join(text_parts)
    except Exception as e:
        log.warning(f"pdfplumber 失敗 ({filename}): {e}")
        text = f"[無法解析 PDF: {e}]"
    finally:
        tmp_path.unlink(missing_ok=True)
    return text


def extract_text_from_docx(docx_bytes: bytes, filename: str) -> str:
    """從 DOCX bytes 擷取純文字"""
    import io
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 也嘗試抓表格內容
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))
        return "\n".join(paragraphs)
    except Exception as e:
        log.warning(f"docx 解析失敗 ({filename}): {e}")
        return f"[無法解析 DOCX: {e}]"


def save_text(slug: str, content: str, meta: dict) -> Path:
    """把內容存成 .txt 並附上 meta 資訊"""
    out_path = OUTPUT_DIR / f"{slug}.txt"
    header = (
        f"來源: {meta.get('url', '')}\n"
        f"標題: {meta.get('title', '')}\n"
        f"爬取時間: {meta.get('scraped_at', '')}\n"
        f"標籤: {meta.get('tag', '')}\n"
        f"{'='*60}\n\n"
    )
    out_path.write_text(header + content, encoding="utf-8")
    return out_path


# ── 核心爬蟲邏輯 ──────────────────────────────────

def scrape_page(page_cfg: dict) -> list[dict]:
    """爬取單一頁面，回傳所有找到的文件資訊"""
    url = page_cfg["url"]
    tag = page_cfg["tag"]
    name = page_cfg["name"]
    log.info(f"▶ 爬取: {name} ({url})")

    resp = safe_get(url)
    if resp is None:
        log.error(f"  ✗ 無法連線: {url}")
        return []

    soup = BeautifulSoup(resp.content, "lxml")
    results = []

    # 1. 儲存頁面本身的文字
    page_text = extract_page_text(soup)
    if page_text.strip():
        slug = url_to_filename(url)
        meta = {
            "url": url,
            "title": name,
            "scraped_at": datetime.now().isoformat(),
            "tag": tag,
            "type": "webpage",
        }
        out = save_text(slug, page_text, meta)
        results.append({**meta, "output_file": str(out)})
        log.info(f"  ✓ 頁面文字已儲存 → {out.name}")

    # 2. 找出所有連結，遞迴處理 PDF / DOCX / 子頁
    links = find_relevant_links(soup, url)
    for link_url, link_text in links:
        ext = urlparse(link_url).path.lower().split(".")[-1]
        if ext in ("pdf", "doc", "docx"):
            doc_result = download_and_convert_doc(link_url, link_text, tag)
            if doc_result:
                results.append(doc_result)
        elif is_course_related(link_text):
            sub_result = scrape_subpage(link_url, link_text, tag)
            if sub_result:
                results.append(sub_result)

    return results


def extract_page_text(soup: BeautifulSoup) -> str:
    """從 HTML 擷取乾淨的可讀文字"""
    # 移除 script / style / nav / footer
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    # 嘗試找主內容區塊
    main = (
        soup.find("main")
        or soup.find("article")
        or soup.find("div", class_=re.compile(r"content|main|post|entry", re.I))
        or soup.find("div", id=re.compile(r"content|main|post|entry", re.I))
        or soup.body
    )

    if main is None:
        return soup.get_text(separator="\n", strip=True)

    lines = []
    for elem in main.descendants:
        if elem.name in ("h1", "h2", "h3", "h4"):
            lines.append(f"\n## {elem.get_text(strip=True)}")
        elif elem.name in ("p", "li", "td", "th"):
            t = elem.get_text(strip=True)
            if t:
                lines.append(t)

    return "\n".join(lines)


def find_relevant_links(soup: BeautifulSoup, base_url: str) -> list[tuple[str, str]]:
    """找出頁面中與選課相關的連結"""
    seen = set()
    result = []
    base_domain = urlparse(base_url).netloc

    for a in soup.find_all("a", href=True):
        href = a["href"].strip()
        text = a.get_text(strip=True)
        full_url = urljoin(base_url, href)
        parsed = urlparse(full_url)

        # 跳過錨點、javascript、mail
        if parsed.scheme not in ("http", "https"):
            continue
        if full_url in seen:
            continue
        seen.add(full_url)

        ext = parsed.path.lower().split(".")[-1]

        # 強制納入 PDF / DOC
        if ext in ("pdf", "doc", "docx"):
            result.append((full_url, text or href))
            continue

        # 只跟蹤 mcu.edu.tw 的子連結
        if "mcu.edu.tw" not in parsed.netloc:
            continue

        # 判斷是否與選課相關
        if is_course_related(text) or is_course_related(href):
            result.append((full_url, text or href))

    return result


def download_and_convert_doc(url: str, title: str, tag: str) -> dict | None:
    """下載 PDF / DOCX 並轉換成文字"""
    log.info(f"  ↳ 下載文件: {title[:40]} ({url})")
    resp = safe_get(url)
    if resp is None:
        return None

    ext = urlparse(url).path.lower().split(".")[-1]
    filename = url_to_filename(url)

    if ext == "pdf":
        text = extract_text_from_pdf(resp.content, filename)
    elif ext in ("doc", "docx"):
        text = extract_text_from_docx(resp.content, filename)
    else:
        return None

    if not text.strip():
        log.warning(f"    ✗ 文件為空或無法擷取文字: {url}")
        return None

    meta = {
        "url": url,
        "title": title,
        "scraped_at": datetime.now().isoformat(),
        "tag": tag,
        "type": ext,
    }
    out = save_text(filename, text, meta)
    log.info(f"    ✓ 已轉換 → {out.name} ({len(text)} 字元)")
    return {**meta, "output_file": str(out)}


def scrape_subpage(url: str, title: str, tag: str) -> dict | None:
    """爬取單一子頁面（不再遞迴）"""
    log.info(f"  ↳ 子頁面: {title[:40]} ({url})")
    resp = safe_get(url)
    if resp is None:
        return None

    soup = BeautifulSoup(resp.content, "lxml")
    text = extract_page_text(soup)
    if not text.strip():
        return None

    # 同時也在子頁找文件
    links = find_relevant_links(soup, url)
    extra_texts = []
    for link_url, link_text in links:
        ext = urlparse(link_url).path.lower().split(".")[-1]
        if ext in ("pdf", "doc", "docx"):
            r = safe_get(link_url)
            if r:
                if ext == "pdf":
                    doc_text = extract_text_from_pdf(r.content, url_to_filename(link_url))
                else:
                    doc_text = extract_text_from_docx(r.content, url_to_filename(link_url))
                if doc_text.strip():
                    extra_texts.append(
                        f"\n\n[附件: {link_text} — {link_url}]\n{doc_text}"
                    )

    combined = text + "".join(extra_texts)
    meta = {
        "url": url,
        "title": title,
        "scraped_at": datetime.now().isoformat(),
        "tag": tag,
        "type": "webpage",
    }
    slug = url_to_filename(url)
    out = save_text(slug, combined, meta)
    log.info(f"    ✓ 子頁面已儲存 → {out.name}")
    return {**meta, "output_file": str(out)}


# ── 匯總報告 ──────────────────────────────────────

def generate_index(all_results: list[dict]) -> None:
    """產生 index.md 與 index.json 摘要"""
    now = datetime.now().isoformat()

    # JSON index
    index_json = OUTPUT_DIR / "index.json"
    index_data = {
        "generated_at": now,
        "total_documents": len(all_results),
        "documents": all_results,
    }
    index_json.write_text(json.dumps(index_data, ensure_ascii=False, indent=2), encoding="utf-8")

    # Markdown index
    lines = [
        "# 銘傳大學選課資訊索引",
        f"",
        f"> 最後更新：{now}  ",
        f"> 共 {len(all_results)} 份文件",
        "",
        "---",
        "",
    ]

    by_tag: dict[str, list[dict]] = {}
    for r in all_results:
        by_tag.setdefault(r.get("tag", "other"), []).append(r)

    tag_labels = {
        "announcement": "📢 教務處公告",
        "freshman": "🎓 新生選課",
        "english": "🔤 英語教學中心",
        "course_structure": "📚 開課與課程架構",
        "other": "📄 其他",
    }

    for tag, docs in by_tag.items():
        lines.append(f"## {tag_labels.get(tag, tag)}")
        lines.append("")
        for d in docs:
            fname = Path(d["output_file"]).name
            lines.append(f"- [{d['title']}]({fname}) `{d['type']}` — {d['url']}")
        lines.append("")

    (OUTPUT_DIR / "index.md").write_text("\n".join(lines), encoding="utf-8")
    log.info(f"✅ 索引已產生 → output/index.md & output/index.json")


# ── 主程式入口 ────────────────────────────────────

def main():
    log.info("=" * 60)
    log.info("銘傳大學選課資訊爬蟲 啟動")
    log.info(f"時間: {datetime.now().isoformat()}")
    log.info("=" * 60)

    all_results = []
    for page_cfg in TARGET_PAGES:
        try:
            results = scrape_page(page_cfg)
            all_results.extend(results)
        except Exception as e:
            log.error(f"爬取 {page_cfg['name']} 時發生錯誤: {e}", exc_info=True)
        time.sleep(1)  # 禮貌性延遲

    generate_index(all_results)
    log.info(f"\n🎉 完成！共儲存 {len(all_results)} 份文件至 output/")


if __name__ == "__main__":
    main()
