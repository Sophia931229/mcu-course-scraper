"""
銘傳大學選課資訊爬蟲 v3.0
"""

import os
import re
import json
import time
import hashlib
import logging
import requests
import pdfplumber
import docx
from datetime import datetime
from pathlib import Path
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("scraper.log", encoding="utf-8"),
    ],
)
log = logging.getLogger(__name__)

OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "zh-TW,zh;q=0.9,en;q=0.8",
}

SESSION = requests.Session()
SESSION.headers.update(HEADERS)

TARGET_PAGES = [
    {"name": "教務處最新公告", "url": "https://academic.mcu.edu.tw/", "tag": "announcement"},
    {"name": "教務處公告列表", "url": "https://academic.mcu.edu.tw/?cat=2", "tag": "announcement"},
    {"name": "新生選課注意事項", "url": "https://freshman.mcu.edu.tw/a01/a01-04/", "tag": "freshman"},
    {"name": "英語教學中心選課", "url": "https://elc.mcu.edu.tw/%E9%81%B8%E8%AA%B2/", "tag": "english"},
    {"name": "開課與師資資訊", "url": "https://academic.mcu.edu.tw/coursestructure/", "tag": "course_structure"},
]

COURSE_KEYWORDS = [
    "選課", "course", "課程", "開課", "停課", "加退選",
    "選課辦法", "選課時間", "選課注意", "課表", "學分",
    "暑修", "補修", "重修", "校際", "網路選課", "停班",
]

RECLASSIFY_RULES = [
    (r"加退選|加選|退選|停修|Withdraw",     "add_drop"),
    (r"跨校|校際|優久|聯盟跨",             "cross_school"),
    (r"AI|人工智慧|TAICA|ai_alliance",      "ai_alliance"),
    (r"輔系|雙主修|學分學程|eForm",         "double_major"),
    (r"暑修|補修|重修|make.?up|summer",     "remedial"),
    (r"停課|停班|補課|調課|cancel",         "class_change"),
    (r"新生|freshman|a01-04",              "freshman"),
    (r"elc\.mcu|英語教學|ELC|english",     "english"),
]

TAG_LABELS = {
    "announcement":     "📢 教務處公告",
    "add_drop":         "🔄 加退選公告",
    "cross_school":     "🏫 校際選課",
    "ai_alliance":      "🤖 AI 聯盟課程",
    "double_major":     "🎓 輔系雙主修",
    "remedial":         "📝 暑修補修重修",
    "class_change":     "⚠️ 停課調課補課",
    "freshman":         "🌱 新生選課說明",
    "english":          "🌐 英語教學中心",
    "course_structure": "📚 開課與課程架構",
    "other":            "📄 其他",
}


def detect_semester(text: str) -> str:
    m = re.search(r"(\d{3,4})學年度?第?(\d)學期", text)
    if m:
        return f"{m.group(1)}-{m.group(2)}"
    return ""


def refine_tag(tag: str, title: str, url: str) -> str:
    if tag != "announcement":
        return tag
    text = (title + " " + url).lower()
    for pattern, new_tag in RECLASSIFY_RULES:
        if re.search(pattern, text, re.IGNORECASE):
            return new_tag
    return tag


def safe_get(url: str, timeout: int = 20):
    for attempt in range(3):
        try:
            resp = SESSION.get(url, timeout=timeout, allow_redirects=True)
            resp.raise_for_status()
            return resp
        except requests.RequestException as e:
            log.warning(f"[嘗試 {attempt+1}/3] 無法取得 {url}: {e}")
            time.sleep(2 ** attempt)
    return None


def url_to_filename(url: str) -> str:
    h = hashlib.md5(url.encode()).hexdigest()[:8]
    parsed = urlparse(url)
    name = re.sub(r"[^\w\-]", "_", parsed.path)[-40:]
    return f"{name}_{h}"


def is_course_related(text: str) -> bool:
    text_lower = text.lower()
    return any(kw in text_lower for kw in COURSE_KEYWORDS)


def extract_text_from_pdf(pdf_bytes: bytes, filename: str) -> str:
    tmp_path = OUTPUT_DIR / f"_tmp_{filename}.pdf"
    tmp_path.write_bytes(pdf_bytes)
    text_parts = []
    try:
        with pdfplumber.open(str(tmp_path)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(f"--- 第 {i} 頁 ---\n{page_text}")
        text = "\n\n".join(text_parts)
    except Exception as e:
        log.warning(f"pdfplumber 失敗 ({filename}): {e}")
        text = f"[無法解析 PDF: {e}]"
    finally:
        tmp_path.unlink(missing_ok=True)
    return text


def extract_text_from_docx(docx_bytes: bytes, filename: str) -> str:
    import io
    try:
        d = docx.Document(io.BytesIO(docx_bytes))
        paragraphs = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))
        return "\n".join(paragraphs)
    except Exception as e:
        log.warning(f"docx 解析失敗 ({filename}): {e}")
        return f"[無法解析 DOCX: {e}]"


def save_text(slug: str, content: str, meta: dict) -> Path:
    out_path = OUTPUT_DIR / f"{slug}.txt"
    header = (
        f"來源: {meta.get('url', '')}\n"
        f"標題: {meta.get('title', '')}\n"
        f"爬取時間: {meta.get('scraped_at', '')}\n"
        f"標籤: {meta.get('tag', '')}\n"
        f"{'='*60}\n\n"
    )
    out_path.write_text(header + content, encoding="utf-8")
    return out_path


def extract_page_text(soup: BeautifulSoup) -> str:
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()
    main = (
        soup.find("main")
        or soup.find("article")
        or soup.find("div", class_=re.compile(r"content|main|post|entry", re.I))
        or soup.find("div", id=re.compile(r"content|main|post|entry", re.I))
        or soup.body
    )
    if main is None:
        return soup.get_text(separator="\n", strip=True)
    lines = []
    for elem in main.descendants:
        if elem.name in ("h1", "h2", "h3", "h4"):
            lines.append(f"\n## {elem.get_text(strip=True)}")
        elif elem.name in ("p", "li", "td", "th"):
            t = elem.get_text(strip=True)
            if t:
                lines.append(t)
    return "\n".join(lines)


def find_relevant_links(soup: BeautifulSoup, base_url: str) -> list[tuple[str, str]]:
    seen = set()
    result = []
    for a in soup.find_all("a", href=True):
        href = a["href"].strip()
        text = a.get_text(strip=True)
        full_url = urljoin(base_url, href)
        parsed = urlparse(full_url)
        if parsed.scheme not in ("http", "https"):
            continue
        if full_url in seen:
            continue
        seen.add(full_url)
        ext = parsed.path.lower().split(".")[-1]
        if ext in ("pdf", "doc", "docx"):
            result.append((full_url, text or href))
            continue
        if "mcu.edu.tw" not in parsed.netloc:
            continue
        if is_course_related(text) or is_course_related(href):
            result.append((full_url, text or href))
    return result


def scrape_page(page_cfg: dict) -> list[dict]:
    url  = page_cfg["url"]
    tag  = page_cfg["tag"]
    name = page_cfg["name"]
    log.info(f"▶ 爬取: {name} ({url})")

    resp = safe_get(url)
    if resp is None:
        log.error(f"  ✗ 無法連線: {url}")
        return []

    soup = BeautifulSoup(resp.content, "lxml")
    results = []

    page_text = extract_page_text(soup)
    if page_text.strip():
        slug = url_to_filename(url)
        final_tag = refine_tag(tag, name, url)
        meta = {
            "url": url,
            "title": name,
            "scraped_at": datetime.now().isoformat(),
            "tag": final_tag,
            "type": "webpage",
            "semester": detect_semester(name + " " + url),
            "output_file": f"output/{slug}.txt",
        }
        save_text(slug, page_text, meta)
        results.append(meta)
        log.info(f"  ✓ 頁面已儲存 [tag={final_tag}]")

    links = find_relevant_links(soup, url)
    for link_url, link_text in links:
        ext = urlparse(link_url).path.lower().split(".")[-1]
        if ext in ("pdf", "doc", "docx"):
            doc_result = download_and_convert_doc(link_url, link_text, tag)
            if doc_result:
                results.append(doc_result)
        elif is_course_related(link_text):
            sub_result = scrape_subpage(link_url, link_text, tag)
            if sub_result:
                results.append(sub_result)

    return results


def download_and_convert_doc(url: str, title: str, base_tag: str) -> dict | None:
    log.info(f"  ↳ 下載文件: {title[:40]} ({url})")
    resp = safe_get(url)
    if resp is None:
        return None
    ext = urlparse(url).path.lower().split(".")[-1]
    filename = url_to_filename(url)
    if ext == "pdf":
        text = extract_text_from_pdf(resp.content, filename)
    elif ext in ("doc", "docx"):
        text = extract_text_from_docx(resp.content, filename)
    else:
        return None
    if not text.strip():
        log.warning(f"    ✗ 文件為空: {url}")
        return None
    final_tag = refine_tag(base_tag, title, url)
    meta = {
        "url": url,
        "title": title,
        "scraped_at": datetime.now().isoformat(),
        "tag": final_tag,
        "type": ext,
        "semester": detect_semester(title + " " + url),
        "output_file": f"output/{filename}.txt",
    }
    save_text(filename, text, meta)
    log.info(f"    ✓ 已轉換 → {filename}.txt [tag={final_tag}]")
    return meta


def scrape_subpage(url: str, title: str, base_tag: str) -> dict | None:
    log.info(f"  ↳ 子頁面: {title[:40]} ({url})")
    resp = safe_get(url)
    if resp is None:
        return None
    soup = BeautifulSoup(resp.content, "lxml")
    text = extract_page_text(soup)
    if not text.strip():
        return None
    links = find_relevant_links(soup, url)
    extra_texts = []
    for link_url, link_text in links:
        ext = urlparse(link_url).path.lower().split(".")[-1]
        if ext in ("pdf", "doc", "docx"):
            r = safe_get(link_url)
            if r:
                doc_text = (
                    extract_text_from_pdf(r.content, url_to_filename(link_url))
                    if ext == "pdf"
                    else extract_text_from_docx(r.content, url_to_filename(link_url))
                )
                if doc_text.strip():
                    extra_texts.append(f"\n\n[附件: {link_text} — {link_url}]\n{doc_text}")
    combined = text + "".join(extra_texts)
    final_tag = refine_tag(base_tag, title, url)
    slug = url_to_filename(url)
    meta = {
        "url": url,
        "title": title,
        "scraped_at": datetime.now().isoformat(),
        "tag": final_tag,
        "type": "webpage",
        "semester": detect_semester(title + " " + url),
        "output_file": f"output/{slug}.txt",
    }
    save_text(slug, combined, meta)
    log.info(f"    ✓ 子頁面已儲存 → {slug}.txt [tag={final_tag}]")
    return meta


def generate_index(all_results: list[dict]) -> None:
    now = datetime.now().isoformat()
    seen_urls: dict[str, dict] = {}
    for r in all_results:
        url = r.get("url", "")
        if url not in seen_urls or r["scraped_at"] > seen_urls[url]["scraped_at"]:
            seen_urls[url] = r
    deduped = list(seen_urls.values())

    tag_counts: dict[str, int] = {}
    for r in deduped:
        t = r.get("tag", "other")
        tag_counts[t] = tag_counts.get(t, 0) + 1

    index_data = {
        "generated_at": now,
        "total_documents": len(deduped),
        "tag_counts": tag_counts,
        "documents": deduped,
    }
    (OUTPUT_DIR / "index.json").write_text(
        json.dumps(index_data, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    lines = [
        "# 銘傳大學選課資訊索引", "",
        f"> 最後更新：{now}  ",
        f"> 共 {len(deduped)} 份文件", "", "---", "",
    ]
    by_tag: dict[str, list[dict]] = {}
    for r in deduped:
        by_tag.setdefault(r.get("tag", "other"), []).append(r)

    tag_order = list(TAG_LABELS.keys())
    for tag, docs in sorted(by_tag.items(), key=lambda x: tag_order.index(x[0]) if x[0] in tag_order else 99):
        lines.append(f"## {TAG_LABELS.get(tag, tag)}")
        lines.append("")
        for d in docs:
            fname = Path(d["output_file"]).name
            lines.append(f"- [{d['title']}]({fname}) `{d['type']}` — {d['url']}")
        lines.append("")

    (OUTPUT_DIR / "index.md").write_text("\n".join(lines), encoding="utf-8")
    log.info("✅ 索引已產生 → output/index.md & output/index.json")
    log.info("📊 分類統計：")
    for t, cnt in sorted(tag_counts.items(), key=lambda x: -x[1]):
        log.info(f"   {TAG_LABELS.get(t, t)}: {cnt} 筆")


def main():
    log.info("=" * 60)
    log.info("銘傳大學選課資訊爬蟲 v3.0 啟動")
    log.info(f"時間: {datetime.now().isoformat()}")
    log.info("=" * 60)

    all_results = []
    for page_cfg in TARGET_PAGES:
        try:
            results = scrape_page(page_cfg)
            all_results.extend(results)
        except Exception as e:
            log.error(f"爬取 {page_cfg['name']} 時發生錯誤: {e}", exc_info=True)
        time.sleep(1)

    generate_index(all_results)
    log.info(f"\n🎉 完成！共儲存 {len(all_results)} 份文件至 output/")


if __name__ == "__main__":
    main()
